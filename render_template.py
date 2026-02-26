from docx import Document
from copy import deepcopy
import json

def replace_text_in_paragraph(paragraph, data):
    full_text = "".join(run.text for run in paragraph.runs)

    replaced = False
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value))
            replaced = True

    if replaced:
        # borrar runs existentes
        for run in paragraph.runs:
            run.text = ""

        # escribir todo en el primer run
        paragraph.runs[0].text = full_text


def replace_all(doc, data):
    # párrafos normales
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, data)

    # tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, data)
doc = Document("tabla_template.docx")

datos = {
    "titulo": "Informe Mensual",
    "cliente": "Claro",
    "fecha": "26/02/2026",
    "empresa": "Mi Empresa SA"
}

replace_all(doc, datos)


tabla = doc.tables[0]

fila_modelo = tabla.rows[1]


# Cargar los datos desde el JSON
with open("datos.json", "r", encoding="utf-8") as f:
    datos = json.load(f)
    
tabla0 = datos["tabla0"]
tabla1 = datos["tabla1"]   

for i, d in enumerate(tabla0):
    if i == 0:
        fila = fila_modelo
    else:
        # agregar fila y copiar celdas de la fila modelo
        fila = tabla.add_row()
        for j, cell in enumerate(fila.cells):
            # copiar el contenido y el estilo de la celda modelo
            new_cell = fila_modelo.cells[j]
            cell._tc.clear_content()  # limpia contenido
            for element in new_cell._tc:
                cell._tc.append(deepcopy(element))

    # ahora ponemos los datos
    fila.cells[0].text = d["nombre"]
    fila.cells[1].text = d["apellido"]
    fila.cells[2].text = d["email"]
    
tabla = doc.tables[1]   
fila_modelo = tabla.rows[1]
    
for i, d in enumerate(tabla1):
    if i == 0:
        fila = fila_modelo
    else:
        # agregar fila y copiar celdas de la fila modelo
        fila = tabla.add_row()
        for j, cell in enumerate(fila.cells):
            # copiar el contenido y el estilo de la celda modelo
            new_cell = fila_modelo.cells[j]
            cell._tc.clear_content()  # limpia contenido
            for element in new_cell._tc:
                cell._tc.append(deepcopy(element))

    # ahora ponemos los datos
    fila.cells[0].text = d["nombre"]
    fila.cells[1].text = d["ip"]
    fila.cells[2].text = d["SO"]
    fila.cells[3].text = d["servicios"]

doc.save("NIP_generado.docx")
